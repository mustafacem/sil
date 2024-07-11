import streamlit as st
from chat_gpt.chat_gpt_c import * #ask_chatgpt, process_items
from doc_creation.doc_creation_c import * #create_document, create_excel_with_values, update_excel_file
from ocr_processing.ocr_processing_c import *
from whisper_speech_to_text.whisper_speech_to_text_c import transcribe_czech_audio


from io import BytesIO
from docx import Document

def create_docx(items_dict):
    doc = Document()
    doc.add_heading('Use Case Description', 0)

    for key, value in items_dict.items():
        doc.add_heading(key, level=1)
        doc.add_paragraph(value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main():
    st.title("Proposal Droid_01")

    openai_api_key = st.text_input("Enter your z API key:", type="password", key="api_key_input")
    if openai_api_key:
        # Initialize OpenAI API with the provided key
        openai.api_key = openai_api_key

    if "notes" not in st.session_state:
        st.session_state.notes = None

    if "items_dict" not in st.session_state:
        st.session_state.items_dict = {
            "Target functionality": None,
            "Solution": None,
            "Inputs": None,
            "Outputs": None,
            "Requirements and assumptions": None,
            "PoC vs production implementation": None,
            "Human review of outputs": None,
            "Post-PoC improvements and functionality add-ons": None,
            "Infrastructure": None,
            "Limitations": None
        }

    if "current_item_index" not in st.session_state:
        st.session_state.current_item_index = 0

    if "notes_finalized" not in st.session_state:
        st.session_state.notes_finalized = False

    if "use_case_finalized" not in st.session_state:
        st.session_state.use_case_finalized = False

    if "use_case_description" not in st.session_state:
        st.session_state.use_case_description = ""

    file_type = st.selectbox("Select the type of notes you want to upload:", ["Image", "Voice Recording"])
    uploaded_file = st.file_uploader("Upload your file", type=["png", "jpg", "jpeg", "wav", "mp3", "m4a"])

    if uploaded_file is not None:
        if file_type == "Image":
            st.image(uploaded_file, caption="Uploaded Image", use_column_width=True)
            preprocess = st.checkbox("Preprocess Image")
            preprocessed_image = uploaded_file

            if preprocess:
                preprocessed_image = preprocess_handwritten_image(uploaded_file)
                if preprocessed_image is not None:
                    st.image(preprocessed_image, caption="Preprocessed Image", use_column_width=True)

            ocr_method = st.selectbox("Select OCR method:", ["Method 1", "Method 2", "Method 3"])
            if st.button("Extract Text from Image"):
                if preprocessed_image is not None:
                    image_to_process = preprocessed_image
                else:
                    image_to_process = uploaded_file

                if ocr_method == "Method 1":
                    st.session_state.notes = extract_text_from_image(openai_api_key, image_to_process)
                elif ocr_method == "Method 2":
                    ocr_output = tes_ext(image_to_process)
                    st.session_state.notes = ocr_output
                elif ocr_method == "Method 3":
                    # Add your method 3 implementation here
                    pass

            if 'notes' in st.session_state:
                use_decoder = st.checkbox("Use decoder prompt?")
                if use_decoder:
                    decoder_prompt = st.selectbox("Select decoder prompt:", ["Prompt 1", "Prompt 2", "Prompt 3", "Prompt 4"])
                    if st.button("Apply Decoder Prompt"):
                        if decoder_prompt == "Prompt 1":
                            decoder_prompt = decoder_prompt_1
                        elif decoder_prompt == "Prompt 2":
                            decoder_prompt = decoder_prompt_2
                        elif decoder_prompt == "Prompt 3":
                            decoder_prompt = decoder_prompt_3
                        elif decoder_prompt == "Prompt 4":
                            decoder_prompt = decoder_prompt_4
                        st.session_state.notes = decoder(st.session_state.notes, decoder_prompt)
                        st.write(st.session_state.notes)
                else:
                    st.write(st.session_state.notes)
                    # Uncomment and define ocr_method_3 if available
                    # elif ocr_method == "Method 3":
                    #     st.session_state.notes = ocr_method_3(image_to_process)

        elif file_type == "Voice Recording":
            st.audio(uploaded_file, format="audio/wav" if uploaded_file.name.endswith('.wav') else "audio/mp3")
            language = st.radio("Select the language of the audio:", ("Czech", "English"))

            if st.button("Transcribe Audio"):
                if language == "Czech":
                    with st.spinner('Transcribing audio...'):
                        st.session_state.notes = transcribe_czech_audio(uploaded_file)
                elif language == "English":
                    with st.spinner('Transcribing audio...'):
                        st.session_state.notes = transcribe_english_audio(uploaded_file)

                st.success("Transcription completed!")

    if st.session_state.notes:
        st.write("Transcribed Notes:")
        st.write(st.session_state.notes)

        if st.button("Edit Notes"):
            st.session_state.edit_mode = True
        if "edit_mode" in st.session_state and st.session_state.edit_mode:
            edited_notes = st.text_area("Edit your notes here:", st.session_state.notes, key="edit_notes_input")
            if st.button("Save Edited Notes"):
                st.session_state.notes = edited_notes
                st.session_state.edit_mode = False
                st.success("Notes updated!")
        if st.button("Use Notes As Is"):
            st.session_state.notes_finalized = True
            st.success("Using transcribed notes as is.")

    if st.session_state.notes_finalized and not st.session_state.use_case_description:
        st.session_state.use_case_description = ask_chatgpt("generate very short use case description from given notes: " + st.session_state.notes)

    if st.session_state.use_case_description:
        st.write("Use Case Description:")
        st.write(st.session_state.use_case_description)
        st.session_state.use_case_description = st.text_area("Edit Use Case Description:", st.session_state.use_case_description, key="use_case_desc_input")
        if st.button("Finalize Use Case Description"):
            st.session_state.use_case_finalized = True

    if st.session_state.use_case_finalized:
        keys = list(st.session_state.items_dict.keys())

        if st.session_state.current_item_index < len(keys):
            current_key = keys[st.session_state.current_item_index]

            if st.session_state.items_dict[current_key] is None:
                prompt = f" generate a short 2-4 sentence {current_key} for business proposal from : {st.session_state.notes} "
                value = ask_chatgpt(prompt)
                st.session_state.items_dict[current_key] = value
                opion_ai = checker(st.session_state.items_dict[current_key], value)
            else:
                opion_ai = ""

            st.write(f"{current_key}: {st.session_state.items_dict[current_key]}")
            st.write(f" AI Opinion for {current_key}: \n {opion_ai}")

            action = st.selectbox(f"Select action for {current_key}:", ["Accept", "Edit", "Add"], key=f"action_select_{current_key}")

            if action == "Accept" and st.button(f"Confirm {current_key}", key=f"confirm_{current_key}"):
                st.session_state.current_item_index += 1
                st.success(f"{current_key} accepted!")

            elif action == "Edit":
                edited_value = st.text_input(f"Edit {current_key}:", st.session_state.items_dict[current_key], key=f"edit_{current_key}")
                if st.button(f"Save {current_key}", key=f"save_{current_key}"):
                    st.session_state.items_dict[current_key] = edited_value
                    st.session_state.current_item_index += 1
                    st.success(f"{current_key} updated!")
                    opion_ai = checker(st.session_state.items_dict[current_key], edited_value)

            elif action == "Add":
                additional_value = st.text_input(f"Add to {current_key}:", key=f"add_{current_key}")
                if st.button(f"Add to {current_key}", key=f"add_confirm_{current_key}"):
                    prompt_for_reg = f"{additional_value} generate {st.session_state.items_dict[current_key]} for business proposal for {st.session_state.use_case_description} from : {st.session_state.notes} "
                    st.session_state.items_dict[current_key] = ask_chatgpt(prompt_for_reg)
                    st.session_state.current_item_index += 1
                    st.success(f"Added to {current_key}!")

        if st.session_state.current_item_index >= len(keys):
            st.success("All items processed!")
            docx_buffer = create_document_2("Screenshot 2024-05-23 171557.png", st.session_state.items_dict)
            st.download_button(
                label="Download DOCX",
                data=docx_buffer,
                file_name='output.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # create_excel_with_values_2(st.session_state.notes)
            excel_buffer = create_excel_with_values_2(st.session_state.notes)
            st.download_button(
                label="Download Excel",
                data=excel_buffer,
                file_name='price_estimate.xlsx',
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    if st.button("Reset"):
        st.session_state.clear()
        st.success("Reset completed!")



if __name__ == "__main__":
    main()


